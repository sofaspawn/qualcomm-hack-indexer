"""Document extraction layer for the Lore semantic memory system.

Provides a unified API for extracting structured text from multiple
document formats. Currently supports TXT, PDF, and DOCX. New formats
can be added by implementing a format extractor and registering it in
the extractor registry.

Usage:
    from pc.extractor import extract

    doc = extract("research_paper.pdf")
    print(doc.text[:500])
"""

from __future__ import annotations

import logging
import time
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

import docx
import fitz

from pc.exceptions import (
    CorruptedDocumentError,
    ExtractorError,
    FileNotReadableError,
    UnsupportedFileTypeError,
)
from pc.models import Document

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Abstract base extractor
# ---------------------------------------------------------------------------


class BaseExtractor(ABC):
    """Abstract base class for format-specific extractors.

    Subclasses must implement ``extract_text`` and ``extract_metadata``
    to handle a specific file format.
    """

    @abstractmethod
    def extract_text(self, path: Path) -> tuple[str, int | None]:
        """Extract the full text content from the file.

        Args:
            path: Resolved path to the source file.

        Returns:
            A tuple of (extracted_text, page_count). ``page_count`` is
            ``None`` for formats where the concept does not apply.
        """

    @abstractmethod
    def extract_metadata(self, path: Path) -> dict[str, Any]:
        """Extract format-specific metadata from the file.

        Args:
            path: Resolved path to the source file.

        Returns:
            A dictionary of metadata key-value pairs.
        """


# ---------------------------------------------------------------------------
# TXT extractor
# ---------------------------------------------------------------------------


class TXTExtractor(BaseExtractor):
    """Extractor for plain UTF-8 text files."""

    def extract_text(self, path: Path) -> tuple[str, int | None]:
        """Read a plain text file preserving line breaks and formatting.

        Strips trailing null bytes that some editors embed. Raises a
        ``CorruptedDocumentError`` if the file cannot be decoded as UTF-8.

        Args:
            path: Resolved path to the text file.

        Returns:
            Tuple of (text_content, None). Page count is not applicable.
        """
        try:
            raw = path.read_bytes()
        except OSError as exc:
            raise FileNotReadableError(
                f"Could not read file: {exc}", path=str(path)
            ) from exc

        # Strip trailing null bytes.
        raw = raw.rstrip(b"\x00")

        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise CorruptedDocumentError(
                f"UTF-8 decoding failed for '{path.name}': {exc}",
                path=str(path),
                original_error=exc,
            ) from exc

        return text, None

    def extract_metadata(self, path: Path) -> dict[str, Any]:
        """Return minimal metadata for a text file.

        Args:
            path: Resolved path to the text file.

        Returns:
            Dictionary with encoding information.
        """
        return {"encoding": "utf-8"}


# ---------------------------------------------------------------------------
# PDF extractor
# ---------------------------------------------------------------------------


class PDFExtractor(BaseExtractor):
    """Extractor for PDF documents using PyMuPDF (fitz)."""

    def extract_text(self, path: Path) -> tuple[str, int]:
        """Extract text from every page of a PDF document.

        Pages are concatenated in order with double-newline separators
        to preserve readable spacing.

        Args:
            path: Resolved path to the PDF file.

        Returns:
            Tuple of (full_text, page_count).

        Raises:
            CorruptedDocumentError: If PyMuPDF cannot open or parse the file.
        """
        try:
            pdf = fitz.open(str(path))
        except Exception as exc:
            raise CorruptedDocumentError(
                f"Failed to open PDF '{path.name}': {exc}",
                path=str(path),
                original_error=exc,
            ) from exc

        try:
            pages: list[str] = []
            for page in pdf:
                page_text = page.get_text("text")
                pages.append(page_text)
            page_count = len(pages)
            text = "\n\n".join(pages)
        finally:
            pdf.close()

        return text, page_count

    def extract_metadata(self, path: Path) -> dict[str, Any]:
        """Extract PDF-level metadata (title, author, etc.).

        Args:
            path: Resolved path to the PDF file.

        Returns:
            Dictionary of PDF metadata fields. Empty values are omitted.
        """
        try:
            pdf = fitz.open(str(path))
        except Exception:
            return {}

        try:
            raw_meta: dict[str, Any] = pdf.metadata or {}
            # Filter out empty/None values for cleanliness.
            metadata = {k: v for k, v in raw_meta.items() if v}
        finally:
            pdf.close()

        return metadata


# ---------------------------------------------------------------------------
# DOCX extractor
# ---------------------------------------------------------------------------


class DOCXExtractor(BaseExtractor):
    """Extractor for Microsoft Word (.docx) documents using python-docx."""

    def extract_text(self, path: Path) -> tuple[str, None]:
        """Extract paragraph text from a DOCX document.

        Paragraphs are joined with newline separators to preserve
        the document's structural separation.

        Args:
            path: Resolved path to the DOCX file.

        Returns:
            Tuple of (full_text, None). DOCX does not have a meaningful
            page count without rendering.

        Raises:
            CorruptedDocumentError: If python-docx cannot open the file.
        """
        try:
            doc = docx.Document(str(path))
        except Exception as exc:
            raise CorruptedDocumentError(
                f"Failed to open DOCX '{path.name}': {exc}",
                path=str(path),
                original_error=exc,
            ) from exc

        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n".join(paragraphs)

        return text, None

    def extract_metadata(self, path: Path) -> dict[str, Any]:
        """Extract DOCX core properties (author, title, dates, etc.).

        Args:
            path: Resolved path to the DOCX file.

        Returns:
            Dictionary of core property fields. None values are omitted.
        """
        try:
            doc = docx.Document(str(path))
        except Exception:
            return {}

        props = doc.core_properties
        core_fields = [
            "author",
            "category",
            "comments",
            "content_status",
            "created",
            "identifier",
            "keywords",
            "language",
            "last_modified_by",
            "last_printed",
            "modified",
            "revision",
            "subject",
            "title",
            "version",
        ]

        metadata: dict[str, Any] = {}
        for field_name in core_fields:
            value = getattr(props, field_name, None)
            if value is not None:
                # Convert datetimes to ISO strings for serializability.
                metadata[field_name] = (
                    value.isoformat() if hasattr(value, "isoformat") else value
                )

        return metadata


# ---------------------------------------------------------------------------
# Extractor registry
# ---------------------------------------------------------------------------

_EXTRACTOR_REGISTRY: dict[str, BaseExtractor] = {
    "txt": TXTExtractor(),
    "pdf": PDFExtractor(),
    "docx": DOCXExtractor(),
}


def register_extractor(extension: str, extractor: BaseExtractor) -> None:
    """Register a new format extractor at runtime.

    Allows extending the extraction layer with new formats without
    modifying existing code.

    Args:
        extension: Lowercase file extension without the dot (e.g. "md").
        extractor: An instance of a ``BaseExtractor`` subclass.
    """
    _EXTRACTOR_REGISTRY[extension.lower().lstrip(".")] = extractor
    logger.info("Registered extractor for '.%s'", extension)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def extract(file_path: str | Path) -> Document:
    """Extract structured text and metadata from a document file.

    This is the sole public entry point for the extraction layer. It
    resolves the file path, selects the appropriate format-specific
    extractor, and returns a ``Document`` containing the extracted
    content and metadata.

    Args:
        file_path: Path to the file to extract (string or ``Path``).

    Returns:
        A ``Document`` dataclass with extracted text and metadata.

    Raises:
        FileNotReadableError: If the file does not exist or is not readable.
        UnsupportedFileTypeError: If the file extension has no registered extractor.
        CorruptedDocumentError: If the file cannot be parsed by the underlying library.
        ExtractorError: Base class for any extraction-layer error.

    Example:
        >>> doc = extract("research_paper.pdf")
        >>> print(doc.filename)
        research_paper.pdf
        >>> print(doc.page_count)
        18
    """
    path = Path(file_path).resolve()

    # --- Validate file existence ---
    if not path.exists():
        raise FileNotReadableError(
            f"File not found: '{path}'", path=str(path)
        )
    if not path.is_file():
        raise FileNotReadableError(
            f"Path is not a file: '{path}'", path=str(path)
        )

    extension = path.suffix.lower().lstrip(".")
    if not extension:
        raise UnsupportedFileTypeError(
            extension="(no extension)", path=str(path)
        )

    # --- Select extractor ---
    extractor = _EXTRACTOR_REGISTRY.get(extension)
    if extractor is None:
        raise UnsupportedFileTypeError(extension=extension, path=str(path))

    logger.info("Extraction started: '%s' (format: %s)", path.name, extension)
    start_time = time.perf_counter()

    try:
        text, page_count = extractor.extract_text(path)
        metadata = extractor.extract_metadata(path)
    except ExtractorError:
        # Re-raise our own exceptions as-is.
        raise
    except Exception as exc:
        raise CorruptedDocumentError(
            f"Unexpected error extracting '{path.name}': {exc}",
            path=str(path),
            original_error=exc,
        ) from exc

    elapsed = time.perf_counter() - start_time
    logger.info(
        "Extraction completed: '%s' — %d chars, %.3fs",
        path.name,
        len(text),
        elapsed,
    )

    return Document(
        path=path,
        filename=path.name,
        extension=extension,
        text=text,
        page_count=page_count,
        metadata=metadata,
    )
